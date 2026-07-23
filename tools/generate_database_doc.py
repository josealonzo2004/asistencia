from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Documentacion_Base_Datos_Asiste_U.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F4D78")
        cell.width = widths[index]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            cells[index].width = widths[index]

    doc.add_paragraph()
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Base de Datos del Sistema Asiste U")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Control de Asistencia Universitaria con Laravel, PostgreSQL, QR y GPS")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    add_heading(doc, "1. Objetivo de la base de datos", 1)
    doc.add_paragraph(
        "La base de datos guarda toda la información necesaria para que el sistema controle usuarios, "
        "materias, asignaciones y asistencias. En este proyecto se usa PostgreSQL como motor de base de datos "
        "y Laravel como backend encargado de validar, consultar y guardar la información."
    )

    add_heading(doc, "2. Tablas principales", 1)
    add_table(
        doc,
        ["Tabla", "Qué guarda", "Quién la usa"],
        [
            ["users", "Administradores, docentes y estudiantes.", "Login, admin, docente y estudiante."],
            ["courses", "Materias, horarios, aulas y docente asignado.", "Admin y docente."],
            ["enrollments", "Relación entre estudiantes y materias.", "Admin y validación de asistencia."],
            ["attendance_sessions", "Sesiones de asistencia creadas por el docente.", "Docente y backend."],
            ["attendance_records", "Asistencia individual de cada estudiante.", "Docente, estudiante y reportes."],
            ["personal_access_tokens", "Tokens de sesión generados por Laravel Sanctum.", "Autenticación de la app."],
        ],
        [Inches(1.6), Inches(3.2), Inches(1.7)],
    )

    add_heading(doc, "3. Tabla users", 1)
    doc.add_paragraph(
        "La tabla users guarda todos los usuarios del sistema. No existe una tabla separada para docentes y otra "
        "para estudiantes; se diferencian por el campo role."
    )
    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["id", "Identificador único del usuario."],
            ["name", "Nombre completo."],
            ["email", "Correo institucional, por ejemplo admin@uleam.edu.ec."],
            ["password", "Contraseña encriptada. No se guarda en texto plano."],
            ["role", "Rol del usuario: admin, teacher o student."],
            ["student_code", "Código del estudiante. Solo aplica para usuarios con rol student."],
            ["active", "Indica si el usuario puede iniciar sesión."],
        ],
        [Inches(1.8), Inches(4.7)],
    )

    add_heading(doc, "4. Tabla courses", 1)
    doc.add_paragraph("La tabla courses guarda las materias creadas por el administrador.")
    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["id", "Identificador único de la materia."],
            ["name", "Nombre de la materia."],
            ["code", "Código académico de la materia."],
            ["room", "Aula o laboratorio donde se dicta la clase."],
            ["schedule", "Horario de la materia."],
            ["teacher_id", "Usuario docente asignado a la materia."],
            ["latitude / longitude", "Ubicación base guardada para la materia."],
            ["radius_meters", "Radio permitido para validar asistencia por GPS."],
        ],
        [Inches(1.8), Inches(4.7)],
    )

    add_heading(doc, "5. Tabla enrollments", 1)
    doc.add_paragraph(
        "La tabla enrollments indica qué estudiantes pertenecen a cada materia. Es clave porque el backend "
        "revisa esta tabla antes de permitir que un estudiante marque asistencia."
    )
    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["id", "Identificador de la asignación."],
            ["course_id", "Materia a la que se asigna el estudiante."],
            ["student_id", "Estudiante asignado a la materia."],
        ],
        [Inches(1.8), Inches(4.7)],
    )

    add_heading(doc, "6. Tabla attendance_sessions", 1)
    doc.add_paragraph(
        "Cuando el docente genera un QR, el backend crea una sesión de asistencia en attendance_sessions. "
        "Esta tabla representa la clase abierta para tomar asistencia."
    )
    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["id", "Identificador de la sesión."],
            ["course_id", "Materia relacionada con la sesión."],
            ["teacher_id", "Docente que generó el QR."],
            ["token", "Código interno que viaja dentro del QR."],
            ["starts_at", "Fecha y hora de inicio."],
            ["expires_at", "Fecha y hora de expiración del QR."],
            ["closed_at", "Fecha y hora en que el docente cierra la sesión."],
            ["latitude / longitude", "Ubicación GPS del docente al generar el QR."],
            ["radius_meters", "Distancia máxima permitida entre estudiante y docente."],
        ],
        [Inches(1.8), Inches(4.7)],
    )

    add_heading(doc, "7. Tabla attendance_records", 1)
    doc.add_paragraph(
        "Aquí se guarda la asistencia individual de cada estudiante. Esta es la tabla más importante para "
        "consultar presentes, tardanzas y ausentes."
    )
    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["id", "Identificador del registro de asistencia."],
            ["attendance_session_id", "Sesión de asistencia a la que pertenece el registro."],
            ["student_id", "Estudiante al que pertenece la asistencia."],
            ["status", "Estado: Presente, Tardanza, Ausente o Pendiente."],
            ["latitude / longitude", "Ubicación enviada por el estudiante."],
            ["accuracy", "Precisión del GPS del celular."],
            ["distance_meters", "Distancia calculada entre estudiante y docente."],
            ["validated_at", "Fecha y hora en que se validó la asistencia."],
        ],
        [Inches(1.8), Inches(4.7)],
    )

    add_heading(doc, "8. Tabla personal_access_tokens", 1)
    doc.add_paragraph(
        "Esta tabla es utilizada por Laravel Sanctum. Guarda los tokens que permiten que la aplicación móvil "
        "haga peticiones autenticadas después del login."
    )

    add_heading(doc, "9. Flujo de guardado de asistencia", 1)
    add_bullets(
        doc,
        [
            "El docente inicia sesión y selecciona una materia asignada.",
            "El docente genera un QR; el backend crea un registro en attendance_sessions.",
            "La ubicación GPS del docente queda guardada como referencia de la sesión.",
            "El estudiante escanea el QR desde su app.",
            "La app del estudiante envía el QR, su código estudiantil y su ubicación GPS.",
            "El backend valida QR, horario, materia, matrícula y distancia.",
            "Si todo está correcto, se guarda un registro en attendance_records como Presente.",
            "Cuando el docente cierra la sesión, los estudiantes que no registraron asistencia quedan como Ausente.",
        ],
    )

    add_heading(doc, "10. Resumen final", 1)
    add_table(
        doc,
        ["Información", "Dónde se guarda"],
        [
            ["Usuarios, docentes y estudiantes", "users"],
            ["Materias", "courses"],
            ["Estudiantes asignados a materias", "enrollments"],
            ["QR o sesión creada por el docente", "attendance_sessions"],
            ["Asistencia individual de cada estudiante", "attendance_records"],
            ["Inicio de sesión y tokens", "personal_access_tokens"],
        ],
        [Inches(3.0), Inches(3.5)],
    )

    doc.add_paragraph(
        "En conclusión, attendance_sessions guarda la clase abierta con QR, mientras que attendance_records "
        "guarda el resultado individual de asistencia de cada estudiante."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
